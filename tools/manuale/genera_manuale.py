#!/usr/bin/env python3
"""
Genera il manuale tecnico in formato .docx a partire dai capitoli Markdown
in docs/manuale/.

Uso:
    python tools/manuale/genera_manuale.py

Requisiti:
    pip install python-docx

I capitoli vengono letti in ordine alfabetico di nome file (per questo sono
numerati: 01-, 02-, ...). Ogni file inizia una nuova pagina; il suo titolo H1
diventa il titolo del capitolo.

Il Markdown supportato e' volutamente un sottoinsieme, quello effettivamente
usato nei capitoli:

    # ## ### ####   titoli
    - / * / 1.      liste (due livelli, indentazione a 2 o 4 spazi)
    | a | b |       tabelle GFM (con riga separatrice ---)
    ```lang        blocchi di codice
    > testo         citazione / nota evidenziata
    ---             separatore orizzontale
    **b** *i* `c`   formattazione inline

Tutto il resto viene reso come testo normale.
"""

from __future__ import annotations

import datetime
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:  # pragma: no cover
    sys.exit("Manca python-docx. Installalo con:  pip install python-docx")


# ---------------------------------------------------------------- configurazione

RADICE = Path(__file__).resolve().parents[2]
SORGENTI = RADICE / "docs" / "manuale"
DESTINAZIONE = RADICE / "docs" / "Manuale_Tecnico_Colombini.docx"

TITOLO = "Gestionale Colombini SNC"
SOTTOTITOLO = "Manuale tecnico"
DESTINATARIO = "Documentazione per lo sviluppo e la manutenzione"

# Palette: stesso blu di docs/schema.html e delle stampe PDF del gestionale.
ACCENTO = RGBColor(0x1A, 0x6F, 0xA8)
ACCENTO_HEX = "1A6FA8"
TESTO = RGBColor(0x1F, 0x29, 0x37)
GRIGIO = RGBColor(0x6B, 0x72, 0x80)
CODICE_SFONDO = "F4F4F4"
HEADER_SFONDO = "E8F1F8"
NOTA_SFONDO = "F5F8FB"

FONT_TESTO = "Calibri"
FONT_CODICE = "Consolas"


# ---------------------------------------------------------------- utilita' docx


def _shading(elemento, colore_hex: str) -> None:
    """Applica un colore di sfondo a una cella o a un paragrafo."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), colore_hex)
    elemento.append(shd)


def _bordo_paragrafo(paragrafo, lato: str, colore_hex: str, spessore: int = 12) -> None:
    """Aggiunge un bordo su un lato del paragrafo (usato per le note)."""
    pPr = paragrafo._p.get_or_add_pPr()
    bordi = pPr.find(qn("w:pBdr"))
    if bordi is None:
        bordi = OxmlElement("w:pBdr")
        pPr.append(bordi)
    elemento = OxmlElement(f"w:{lato}")
    elemento.set(qn("w:val"), "single")
    elemento.set(qn("w:sz"), str(spessore))
    elemento.set(qn("w:space"), "8")
    elemento.set(qn("w:color"), colore_hex)
    bordi.append(elemento)


def _campo(paragrafo, istruzione: str, testo_segnaposto: str = "") -> None:
    """
    Inserisce un campo Word (TOC, PAGE, ...) come sequenza di run.
    Word lo calcola all'apertura o alla pressione di F9.
    """
    run = paragrafo.add_run()
    inizio = OxmlElement("w:fldChar")
    inizio.set(qn("w:fldCharType"), "begin")
    run._r.append(inizio)

    run = paragrafo.add_run()
    istr = OxmlElement("w:instrText")
    istr.set(qn("xml:space"), "preserve")
    istr.text = istruzione
    run._r.append(istr)

    run = paragrafo.add_run()
    separatore = OxmlElement("w:fldChar")
    separatore.set(qn("w:fldCharType"), "separate")
    run._r.append(separatore)

    if testo_segnaposto:
        paragrafo.add_run(testo_segnaposto)

    run = paragrafo.add_run()
    fine = OxmlElement("w:fldChar")
    fine.set(qn("w:fldCharType"), "end")
    run._r.append(fine)


def prepara_stili(documento: Document) -> None:
    """Imposta i font e i colori di tutti gli stili usati nel manuale."""
    normale = documento.styles["Normal"]
    normale.font.name = FONT_TESTO
    normale.font.size = Pt(10.5)
    normale.font.color.rgb = TESTO
    normale.paragraph_format.space_after = Pt(7)
    normale.paragraph_format.line_spacing = 1.12

    dimensioni = {"Heading 1": 19, "Heading 2": 14, "Heading 3": 11.5, "Heading 4": 10.5}
    for nome, dimensione in dimensioni.items():
        stile = documento.styles[nome]
        stile.font.name = FONT_TESTO
        stile.font.size = Pt(dimensione)
        stile.font.bold = True
        stile.font.color.rgb = ACCENTO
        stile.paragraph_format.space_before = Pt(16 if nome == "Heading 1" else 12)
        stile.paragraph_format.space_after = Pt(6)
        stile.paragraph_format.keep_with_next = True

    documento.styles["Heading 4"].font.color.rgb = TESTO

    # Gli elenchi non usano gli stili List Bullet/List Number: vedi scrivi_elenco().


def numera_pagine(documento: Document) -> None:
    """Piede di pagina con il numero di pagina centrato."""
    for sezione in documento.sections:
        piede = sezione.footer.paragraphs[0]
        piede.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _campo(piede, " PAGE ", "1")
        for run in piede.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = GRIGIO
            run.font.name = FONT_TESTO


# ---------------------------------------------------------------- testo inline

# Cattura, nell'ordine: `codice`, **grassetto**, *corsivo* o _corsivo_
INLINE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)")


def scrivi_inline(paragrafo, testo: str, grassetto: bool = False, corsivo: bool = False):
    """
    Aggiunge il testo al paragrafo interpretando la formattazione inline.

    Ricorsiva sui marcatori annidati: `**testo con `codice` dentro**` deve
    restare grassetto E monospaziato, non mostrare i backtick.
    """
    for pezzo in INLINE.split(testo):
        if not pezzo:
            continue

        if pezzo.startswith("`") and pezzo.endswith("`"):
            run = paragrafo.add_run(pezzo[1:-1])
            run.font.name = FONT_CODICE
            run.font.size = Pt(9.5)
            run.bold = grassetto
            run.italic = corsivo
        elif pezzo.startswith("**") and pezzo.endswith("**"):
            scrivi_inline(paragrafo, pezzo[2:-2], True, corsivo)
        elif (pezzo.startswith("*") and pezzo.endswith("*")) or (
            pezzo.startswith("_") and pezzo.endswith("_")
        ):
            scrivi_inline(paragrafo, pezzo[1:-1], grassetto, True)
        else:
            run = paragrafo.add_run(pezzo)
            run.bold = grassetto
            run.italic = corsivo

    return paragrafo


# ---------------------------------------------------------------- blocchi


def scrivi_codice(documento: Document, righe: list[str]) -> None:
    """Blocco di codice: monospaziato, sfondo grigio, bordo a sinistra."""
    paragrafo = documento.add_paragraph()
    paragrafo.paragraph_format.left_indent = Cm(0.4)
    paragrafo.paragraph_format.space_before = Pt(4)
    paragrafo.paragraph_format.space_after = Pt(8)
    _shading(paragrafo._p.get_or_add_pPr(), CODICE_SFONDO)
    _bordo_paragrafo(paragrafo, "left", ACCENTO_HEX, 18)

    for indice, riga in enumerate(righe):
        run = paragrafo.add_run(riga)
        run.font.name = FONT_CODICE
        run.font.size = Pt(9)
        if indice < len(righe) - 1:
            run.add_break()


def scrivi_nota(documento: Document, righe: list[str]) -> None:
    """Citazione '>' resa come nota evidenziata con barra laterale."""
    paragrafo = documento.add_paragraph()
    paragrafo.paragraph_format.left_indent = Cm(0.4)
    paragrafo.paragraph_format.space_before = Pt(4)
    paragrafo.paragraph_format.space_after = Pt(8)
    _shading(paragrafo._p.get_or_add_pPr(), NOTA_SFONDO)
    _bordo_paragrafo(paragrafo, "left", ACCENTO_HEX, 18)
    scrivi_inline(paragrafo, " ".join(righe))
    for run in paragrafo.runs:
        run.font.size = Pt(9.5)


def scrivi_tabella(documento: Document, righe: list[str]) -> None:
    """Tabella GFM: prima riga di intestazione, seconda ignorata (separatrice)."""
    def celle(riga: str) -> list[str]:
        return [c.strip() for c in riga.strip().strip("|").split("|")]

    intestazione = celle(righe[0])
    corpo = [celle(r) for r in righe[2:] if r.strip()]
    colonne = len(intestazione)

    tabella = documento.add_table(rows=1, cols=colonne)
    tabella.style = "Table Grid"
    tabella.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabella.autofit = True

    for indice, testo in enumerate(intestazione):
        cella = tabella.rows[0].cells[indice]
        cella.text = ""
        _shading(cella._tc.get_or_add_tcPr(), HEADER_SFONDO)
        paragrafo = cella.paragraphs[0]
        paragrafo.paragraph_format.space_after = Pt(2)
        scrivi_inline(paragrafo, testo, grassetto=True)
        for run in paragrafo.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = ACCENTO

    for valori in corpo:
        cella_riga = tabella.add_row().cells
        for indice in range(colonne):
            testo = valori[indice] if indice < len(valori) else ""
            cella = cella_riga[indice]
            cella.text = ""
            paragrafo = cella.paragraphs[0]
            paragrafo.paragraph_format.space_after = Pt(2)
            scrivi_inline(paragrafo, testo)
            for run in paragrafo.runs:
                run.font.size = Pt(9.5)

    documento.add_paragraph().paragraph_format.space_after = Pt(2)


def scrivi_elenco(documento: Document, riga: str, marcatore: str, livello: int) -> None:
    """
    Voce di elenco costruita a mano invece che con gli stili List Bullet/List Number.

    Quegli stili portano con se' una definizione di numerazione che impone il proprio
    rientro, ignorando sia lo stile sia la formattazione diretta: il risultato e' che
    le righe successive alla prima tornano a filo di margine. Qui il segno di elenco
    e' un run di testo seguito da una tabulazione, con rientro sporgente esplicito:
    la resa e' identica in Word e in qualunque convertitore.
    """
    rientro = 1.5 if livello else 0.75
    paragrafo = documento.add_paragraph()
    formato = paragrafo.paragraph_format
    formato.left_indent = Cm(rientro)
    formato.first_line_indent = Cm(-0.45)
    formato.space_after = Pt(3)
    formato.tab_stops.add_tab_stop(Cm(rientro))

    paragrafo.add_run(f"{marcatore}\t")
    scrivi_inline(paragrafo, riga)


# ---------------------------------------------------------------- parser


def converti_capitolo(documento: Document, testo: str) -> None:
    righe = testo.replace("\r\n", "\n").split("\n")
    indice = 0

    while indice < len(righe):
        riga = righe[indice]
        spoglia = riga.strip()

        # blocco di codice
        if spoglia.startswith("```"):
            indice += 1
            blocco: list[str] = []
            while indice < len(righe) and not righe[indice].strip().startswith("```"):
                blocco.append(righe[indice])
                indice += 1
            indice += 1
            scrivi_codice(documento, blocco or [""])
            continue

        # tabella
        if spoglia.startswith("|") and indice + 1 < len(righe) and set(
            righe[indice + 1].strip()
        ) <= set("|-: "):
            blocco = []
            while indice < len(righe) and righe[indice].strip().startswith("|"):
                blocco.append(righe[indice])
                indice += 1
            if len(blocco) >= 2:
                scrivi_tabella(documento, blocco)
            continue

        # citazione
        if spoglia.startswith(">"):
            blocco = []
            while indice < len(righe) and righe[indice].strip().startswith(">"):
                blocco.append(righe[indice].strip().lstrip(">").strip())
                indice += 1
            scrivi_nota(documento, [r for r in blocco if r])
            continue

        # separatore
        if spoglia in ("---", "***", "___"):
            indice += 1
            continue

        # riga vuota
        if not spoglia:
            indice += 1
            continue

        # titoli
        titolo = re.match(r"^(#{1,4})\s+(.*)$", spoglia)
        if titolo:
            livello = len(titolo.group(1))
            paragrafo = documento.add_heading("", level=livello)
            scrivi_inline(paragrafo, titolo.group(2))
            for run in paragrafo.runs:
                run.font.color.rgb = TESTO if livello >= 4 else ACCENTO
                run.bold = True
            indice += 1
            continue

        # elenchi
        elenco = re.match(r"^(\s*)([-*]|\d+[.)])\s+(.*)$", riga)
        if elenco:
            livello = 1 if len(elenco.group(1)) >= 2 else 0
            simbolo = elenco.group(2)
            marcatore = "•" if simbolo in ("-", "*") else simbolo.replace(")", ".")
            testo = [elenco.group(3)]
            indice += 1

            # Continuazione della voce: righe indentate senza un proprio marcatore.
            # Vanno accorpate, altrimenti diventerebbero paragrafi separati a filo
            # di margine sotto il segno di elenco.
            while indice < len(righe):
                seguente = righe[indice]
                if not seguente.strip() or not seguente.startswith((" ", "\t")):
                    break
                if re.match(r"^\s*([-*]|\d+[.)])\s+", seguente):
                    break
                if re.match(r"^\s*(#{1,4}\s|```|\||>)", seguente):
                    break
                testo.append(seguente.strip())
                indice += 1

            scrivi_elenco(documento, " ".join(testo), marcatore, livello)
            continue

        # paragrafo: accorpa le righe consecutive di testo
        blocco = []
        while indice < len(righe):
            corrente = righe[indice]
            spoglia = corrente.strip()
            if not spoglia or re.match(r"^(#{1,4}\s|```|\||>|\s*([-*]|\d+[.)])\s)", corrente):
                break
            if spoglia in ("---", "***", "___"):
                break
            blocco.append(spoglia)
            indice += 1
        scrivi_inline(documento.add_paragraph(), " ".join(blocco))


# ---------------------------------------------------------------- pagine fisse


def copertina(documento: Document, versione: str) -> None:
    for _ in range(6):
        documento.add_paragraph()

    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run(TITOLO)
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = ACCENTO
    run.font.name = FONT_TESTO

    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run(SOTTOTITOLO)
    run.font.size = Pt(17)
    run.font.color.rgb = TESTO
    run.font.name = FONT_TESTO

    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run(DESTINATARIO)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIGIO
    run.font.italic = True
    run.font.name = FONT_TESTO

    for _ in range(8):
        documento.add_paragraph()

    oggi = datetime.date.today().strftime("%d/%m/%Y")
    for testo in (f"Versione del gestionale: {versione}", f"Documento generato il {oggi}"):
        paragrafo = documento.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragrafo.add_run(testo)
        run.font.size = Pt(10)
        run.font.color.rgb = GRIGIO
        run.font.name = FONT_TESTO

    documento.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def aggiorna_campi_all_apertura(documento: Document) -> None:
    """
    Chiede a Word di ricalcolare i campi (l'indice) all'apertura del documento,
    invece di lasciare all'utente il compito di premere F9.
    """
    impostazioni = documento.settings.element
    if impostazioni.find(qn("w:updateFields")) is None:
        elemento = OxmlElement("w:updateFields")
        elemento.set(qn("w:val"), "true")
        impostazioni.append(elemento)


def indice(documento: Document) -> None:
    paragrafo = documento.add_heading("", level=1)
    scrivi_inline(paragrafo, "Indice")
    for run in paragrafo.runs:
        run.font.color.rgb = ACCENTO
        run.bold = True

    nota = documento.add_paragraph()
    run = nota.add_run(
        "Se le voci non compaiono: in Word, seleziona tutto (Ctrl+A) e premi F9. "
        "LibreOffice non importa questo campo — lì l'indice va inserito con "
        "Inserisci → Sommario e indice."
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GRIGIO

    _campo(documento.add_paragraph(), ' TOC \\o "1-3" \\h \\z \\u ', "")
    documento.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def versione_corrente() -> str:
    """Legge l'ultima versione dichiarata in CHANGELOG.md."""
    changelog = RADICE / "CHANGELOG.md"
    if changelog.exists():
        trovato = re.search(r"^##\s*\[([\d.]+)\]", changelog.read_text(encoding="utf-8"), re.M)
        if trovato:
            return "v" + trovato.group(1)
    return "n/d"


# ---------------------------------------------------------------- main


def main() -> int:
    capitoli = sorted(SORGENTI.glob("*.md"))
    if not capitoli:
        sys.exit(f"Nessun capitolo trovato in {SORGENTI}")

    documento = Document()
    sezione = documento.sections[0]
    sezione.top_margin = Cm(2.2)
    sezione.bottom_margin = Cm(2.2)
    sezione.left_margin = Cm(2.4)
    sezione.right_margin = Cm(2.2)

    prepara_stili(documento)
    documento.core_properties.title = f"{TITOLO} — {SOTTOTITOLO}"
    documento.core_properties.comments = DESTINATARIO

    versione = versione_corrente()
    aggiorna_campi_all_apertura(documento)
    copertina(documento, versione)
    indice(documento)

    for posizione, capitolo in enumerate(capitoli):
        if posizione:
            documento.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        converti_capitolo(documento, capitolo.read_text(encoding="utf-8"))
        print(f"  + {capitolo.name}")

    numera_pagine(documento)
    DESTINAZIONE.parent.mkdir(parents=True, exist_ok=True)
    documento.save(DESTINAZIONE)
    print(f"\nManuale generato: {DESTINAZIONE}  ({versione})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
